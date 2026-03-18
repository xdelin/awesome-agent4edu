#!/usr/bin/env python3
"""Compile draft chapters into a complete manuscript."""

import json
import os
import sys
from datetime import datetime

from update_story_bible import load_bible


def compile_manuscript(path: str, output_format: str = "md") -> str:
    """
    Compile all chapter drafts into a complete manuscript.
    
    Args:
        path: Project directory
        output_format: Output format (md, txt)
    
    Returns:
        Path to compiled manuscript
    """
    bible_path = os.path.join(path, "story-bible.json")
    chapter_dir = os.path.join(path, "draft", "chapters")
    manuscript_dir = os.path.join(path, "manuscript")
    
    os.makedirs(manuscript_dir, exist_ok=True)
    
    # Load story bible using load_bible() for automatic schema validation
    story_bible = load_bible(path)
    
    title = story_bible["meta"]["title"]
    
    # Collect all chapters
    # Filter out example/template keys (starting with _) to avoid ValueError on int()
    chapters = []
    total_words = 0

    chapter_keys = [k for k in story_bible["chapters"].keys() if not k.startswith("_")]
    for chapter_num in sorted([int(k) for k in chapter_keys]):
        chapter_file = os.path.join(chapter_dir, f"ch{chapter_num:02d}.md")
        if os.path.exists(chapter_file):
            with open(chapter_file, "r", encoding="utf-8") as f:
                content = f.read()
            chapters.append({
                "number": chapter_num,
                "content": content,
                "words": story_bible["chapters"][str(chapter_num)]["word_count"]
            })
            total_words += story_bible["chapters"][str(chapter_num)]["word_count"]
    
    # Build manuscript
    manuscript = []
    
    # Title page
    manuscript.append(f"# {title}\n")
    manuscript.append(f"*First Draft*\n")
    manuscript.append(f"*{total_words:,} words*\n")
    manuscript.append(f"*Compiled: {datetime.now().strftime('%Y-%m-%d')}*\n")
    manuscript.append("\n---\n")
    
    # Table of contents
    manuscript.append("\n## Table of Contents\n")
    for ch in chapters:
        ch_title = story_bible["chapters"][str(ch["number"])].get("title", f"Chapter {ch['number']}")
        manuscript.append(f"- Chapter {ch['number']}: {ch_title} ({ch['words']:,} words)")
    manuscript.append("\n---\n")
    
    # Chapters
    for ch in chapters:
        manuscript.append(f"\n{ch['content']}\n")
        manuscript.append("\n---\n")
    
    # Save manuscript
    manuscript_text = "\n".join(manuscript)
    
    output_file = os.path.join(manuscript_dir, f"full-manuscript.{output_format}")
    with open(output_file, "w", encoding="utf-8") as f:
        f.write(manuscript_text)
    
    print(f"[OK] Manuscript compiled: {output_file}")
    print(f"   Chapters: {len(chapters)}")
    print(f"   Total words: {total_words:,}")
    
    # Generate word count report
    report = generate_word_count_report(story_bible, chapters)
    report_file = os.path.join(manuscript_dir, "chapter-word-counts.md")
    with open(report_file, "w", encoding="utf-8") as f:
        f.write(report)
    print(f"   Word count report: {report_file}")
    
    # Generate continuity report
    continuity = generate_continuity_report(story_bible)
    continuity_file = os.path.join(manuscript_dir, "continuity-report.md")
    with open(continuity_file, "w", encoding="utf-8") as f:
        f.write(continuity)
    print(f"   Continuity report: {continuity_file}")
    
    return output_file


def generate_word_count_report(bible: dict, chapters: list) -> str:
    """Generate word count statistics."""
    report = []
    report.append("# Word Count Report\n")
    report.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    
    target = bible["meta"]["target_words"]
    target_per_chapter = bible["meta"]["words_per_chapter"]
    actual_total = bible["progress"]["total_words"]
    
    report.append(f"\n## Summary")
    report.append(f"- Target: {target:,} words")
    report.append(f"- Actual: {actual_total:,} words")
    report.append(f"- Difference: {actual_total - target:+,} words ({(actual_total/target*100):.1f}%)\n")
    
    report.append(f"\n## By Chapter\n")
    report.append("| Chapter | Target | Actual | Diff | Status |")
    report.append("|---------|--------|--------|------|--------|")
    
    for ch in chapters:
        actual = ch["words"]
        diff = actual - target_per_chapter
        pct = actual / target_per_chapter * 100 if target_per_chapter > 0 else 0
        
        if pct < 80:
            status = "[!] Short"
        elif pct > 120:
            status = "[!] Long"
        else:
            status = "[OK]"
        
        report.append(f"| {ch['number']} | {target_per_chapter:,} | {actual:,} | {diff:+,} | {status} |")
    
    return "\n".join(report)


def generate_continuity_report(bible: dict) -> str:
    """Generate a continuity check report."""
    report = []
    report.append("# Continuity Report\n")
    report.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")

    # Unresolved foreshadowing (filter out _comment example entries)
    unresolved = [p for p in bible["foreshadowing"]["planted"]
                  if not p.get("paid_off") and "_comment" not in p]
    report.append(f"\n## Unresolved Foreshadowing ({len(unresolved)} items)\n")
    if unresolved:
        for plant in unresolved:
            report.append(f"- **{plant['thread']}**")
            report.append(f"  - Planted: {plant['planted_in']}")
            if plant.get("expected_payoff"):
                report.append(f"  - Expected: {plant['expected_payoff']}")
    else:
        report.append("All foreshadowing resolved. [OK]\n")
    
    # Resolved foreshadowing (filter out _comment example entries)
    resolved = [p for p in bible["foreshadowing"]["paid_off"] if "_comment" not in p]
    report.append(f"\n## Resolved Foreshadowing ({len(resolved)} items)\n")
    for payoff in resolved:
        report.append(f"- {payoff['thread']}: {payoff['planted_in']} -> {payoff['paid_in']}")
    
    # Invented details (filter out _comment example entries)
    inventions = [d for d in bible["invented_details"] if "_comment" not in d]
    unreviewed = [d for d in inventions if not d.get("reviewed")]
    report.append(f"\n## Invented Details\n")
    report.append(f"- Total: {len(inventions)}")
    report.append(f"- Reviewed: {len(inventions) - len(unreviewed)}")
    report.append(f"- Pending review: {len(unreviewed)}")

    if unreviewed:
        report.append(f"\n### Pending Review:\n")
        for d in unreviewed:
            report.append(f"- [{d['category']}] {d['detail']} (ch{d['chapter']}.{d['scene']})")

    # Established facts count (filter out _comment example entries)
    facts = [f for f in bible["established_facts"] if "_comment" not in f]
    report.append(f"\n## Established Facts: {len(facts)} items\n")

    # Timeline (filter out _comment keys)
    if bible["timeline"]:
        report.append(f"\n## Timeline\n")
        timeline_entries = [(k, v) for k, v in bible["timeline"].items()
                           if not k.startswith("_") and k.isdigit()]
        for ch, time in sorted(timeline_entries, key=lambda x: int(x[0])):
            report.append(f"- Chapter {ch}: {time}")
    
    return "\n".join(report)


def export_for_editing(path: str) -> str:
    """
    Export manuscript in a format suitable for external editing.
    
    Args:
        path: Project directory
    
    Returns:
        Path to exported file
    """
    # Compile basic markdown first
    compile_manuscript(path, "md")
    
    # Try to create docx if python-docx is available
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        
        bible_path = os.path.join(path, "story-bible.json")
        chapter_dir = os.path.join(path, "draft", "chapters")
        manuscript_dir = os.path.join(path, "manuscript")
        
        # Use load_bible for schema validation
        story_bible = load_bible(path)

        title = story_bible["meta"]["title"]

        doc = Document()
        
        # Title
        doc.add_heading(title, 0)
        doc.add_paragraph(f"First Draft — {story_bible['progress']['total_words']:,} words")
        doc.add_page_break()
        
        # Chapters
        chapter_keys = [k for k in story_bible["chapters"].keys() if not k.startswith("_")]
        for chapter_num in sorted([int(k) for k in chapter_keys]):
            chapter_file = os.path.join(chapter_dir, f"ch{chapter_num:02d}.md")
            if os.path.exists(chapter_file):
                with open(chapter_file, "r", encoding="utf-8") as f:
                    content = f.read()

                # Remove markdown headers, add as Word headings
                lines = content.split("\n")
                for line in lines:
                    if line.startswith("# "):
                        doc.add_heading(line[2:], 1)
                    elif line.startswith("## "):
                        doc.add_heading(line[3:], 2)
                    elif line.strip():
                        doc.add_paragraph(line)
                
                doc.add_page_break()
        
        output_file = os.path.join(manuscript_dir, "full-manuscript.docx")
        doc.save(output_file)
        print(f"[OK] Word document created: {output_file}")
        return output_file
        
    except ImportError:
        print("[INFO] python-docx not available. Markdown export only.")
        return os.path.join(path, "manuscript", "full-manuscript.md")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage:")
        print("  compile_manuscript.py <path>")
        print("  compile_manuscript.py <path> --docx")
        sys.exit(1)
    
    path = sys.argv[1]
    
    if len(sys.argv) > 2 and sys.argv[2] == "--docx":
        export_for_editing(path)
    else:
        compile_manuscript(path)
